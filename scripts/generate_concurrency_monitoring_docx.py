from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_PATH = Path(
    "/Users/heybox/Desktop/毕业设计/毕业设计项目version1/shortlink/docs/shortlink-high-concurrency-and-grafana-design.docx"
)
SCREENSHOT_DIR = Path(
    "/Users/heybox/Desktop/毕业设计/毕业设计项目version1/shortlink/docs/assets/screenshots"
)


def setup_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    run.bold = True


def add_bullets(doc: Document, items) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc: Document, image_name: str, caption: str, width: float = 5.8) -> None:
    image_path = SCREENSHOT_DIR / image_name
    if not image_path.exists():
        doc.add_paragraph(f"[缺失图片] {image_name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document()
    setup_style(doc)
    add_title(doc, "短链接系统高并发设计与 Grafana 监控指标设计")

    doc.add_paragraph(
        "本文档从系统实现角度，对短链接系统中的高并发处理方案以及 Grafana 监控指标设计进行整理。"
        "前者关注系统在高访问量条件下如何保证稳定性、正确性和响应效率；后者关注系统上线运行后如何通过监控看板、业务指标与告警规则实现可观测性。"
    )

    doc.add_heading("1. 高并发设计概述", level=1)
    doc.add_paragraph(
        "短链接系统最容易受到高并发冲击的两条核心链路分别是“短链接创建链路”和“短链接跳转链路”。"
        "创建链路主要面临短码冲突、双表写入一致性和缓存同步问题；跳转链路主要面临高并发读、缓存穿透、缓存击穿以及访问配额并发扣减问题。"
        "因此，系统并不是用单一技术处理所有问题，而是按照不同风险点分层设计。"
    )

    doc.add_heading("2. 短链接创建链路的并发设计", level=1)
    doc.add_paragraph(
        "在短链接创建场景中，系统首先会对原始 URL 进行白名单校验，随后使用 MurmurHash32 与 Base62 组合生成短码候选。"
        "由于任何哈希算法都无法绝对避免冲突，因此系统没有把唯一性完全寄托在算法上，而是采用“布隆过滤器预判 + 数据库唯一索引兜底”的双层方案。"
    )
    add_bullets(
        doc,
        [
            "布隆过滤器用于快速预判候选短链是否可能已存在，降低高并发下频繁打数据库的压力。",
            "数据库唯一索引用于在最终写入阶段做强一致裁决，保证数据正确性。",
            "创建逻辑同时写入 t_link 与 t_link_goto 两张表，并通过事务保证双表原子性。",
            "缓存预热发生在事务提交成功之后，避免出现数据库回滚但缓存已写入的幽灵短链问题。",
            "系统保留普通创建和加锁创建两条路径，在吞吐优先与稳定优先之间进行平衡。",
        ],
    )
    doc.add_paragraph(
        "这种设计思路的优点是：常态流量下可以维持较高吞吐量，而在短码冲突压力增大时，仍然能通过唯一约束和加锁策略保证系统正确性。"
    )

    doc.add_heading("3. 短链接跳转链路的并发设计", level=1)
    doc.add_paragraph(
        "跳转链路是系统的核心读路径，其设计目标是尽量让请求在最短路径内完成，不必要时绝不回源数据库。"
        "因此，系统采用了“正向缓存 -> 空值缓存 -> 布隆过滤器 -> 分布式锁保护回源 -> 数据库裁决”的分层决策链。"
    )
    add_bullets(
        doc,
        [
            "优先查询 Redis 正向缓存，命中后直接进入跳转逻辑。",
            "若正向缓存未命中，则查询空值缓存，快速拦截已确认无效的短链请求。",
            "若空值缓存也未命中，则通过布隆过滤器进行存在性预判，减少无效请求回源。",
            "只有 Bloom 可能存在时，才进入分布式锁保护区，避免热点短链缓存失效时的击穿问题。",
            "数据库回源成功后再写回缓存，构成“数据库真相源，缓存派生状态”的一致性原则。",
        ],
    )
    doc.add_paragraph(
        "此外，系统在跳转前还要处理访问配额问题。它没有采用“先查后改”的方式，而是通过单条带条件的 SQL 完成访问次数的原子扣减。"
        "这样即使多个线程同时访问，也不会出现超额消费配额的竞态问题。"
    )

    doc.add_heading("4. 异步统计与高峰期性能优化", level=1)
    doc.add_paragraph(
        "短链接跳转不仅要完成 302 重定向，还要统计 UV、UIP、访问设备、浏览器、地区等信息。"
        "如果这些统计逻辑全部同步执行，会直接抬高用户侧跳转延迟。为此，系统将统计处理尽量异步化，通过 Redis Stream 等方式把统计数据送往后续处理链路，从而把用户感知时间压缩在跳转主流程内。"
    )
    add_bullets(
        doc,
        [
            "用户跳转成功不依赖统计结果落库完成。",
            "统计链路与跳转链路解耦，提升峰值期响应性能。",
            "UV 与 UIP 采用 Cookie、IP 与 Redis 判重组合，形成可接受的运营统计口径。",
        ],
    )

    doc.add_heading("5. 高并发设计总结", level=1)
    doc.add_paragraph(
        "总体来看，系统在高并发下主要解决了五类问题：短码冲突、缓存穿透、缓存击穿、访问配额竞态以及跳转延迟过高。"
        "其核心思想是把问题拆开处理：布隆过滤器与空值缓存负责挡住无效流量，Redis 负责承接热点读，分布式锁负责收敛回源风暴，数据库唯一约束和原子 SQL 负责最终正确性，异步统计负责控制用户感知延迟。"
    )

    doc.add_heading("6. Grafana 监控体系设计概述", level=1)
    doc.add_paragraph(
        "在监控设计方面，系统采用了 Spring Boot Actuator + Micrometer + Prometheus + Grafana + Alertmanager 的组合。"
        "Actuator 暴露标准指标采集端点，Prometheus 负责定时抓取，Grafana 负责可视化展示，Alertmanager 负责告警通知。"
        "这套方案既保留了基础设施指标，也补充了面向短链接业务的语义指标。"
    )
    add_bullets(
        doc,
        [
            "Prometheus 每 15 秒抓取一次聚合服务和 Redis Exporter 指标。",
            "Grafana 看板覆盖总览、业务链路、日志探索等多个维度。",
            "Alert 规则覆盖服务可用性、错误率、延迟、JVM、Redis 与 SLO 错误预算。",
        ],
    )

    doc.add_heading("7. 为什么监控不能只看 CPU 和内存", level=1)
    doc.add_paragraph(
        "如果监控只停留在 CPU、内存、磁盘等基础指标层面，虽然能发现机器层面的异常，但无法直接判断业务是否退化。"
        "例如，系统 CPU 很正常，但短链接创建冲突明显增多；又或者 JVM 并不繁忙，但统计接口空结果率异常升高。"
        "因此，本系统的监控被设计为三层结构：基础资源指标、服务方法指标、业务语义指标。"
    )

    doc.add_heading("8. Grafana 总览指标设计", level=1)
    doc.add_paragraph(
        "总览看板主要回答“系统现在稳不稳定”。因此它重点展示 QPS、P95、5xx 错误率、服务可用性、JVM Heap 使用率、GC 平均停顿、Redis 内存和 Redis 命令吞吐。"
    )
    add_bullets(
        doc,
        [
            "QPS 用于判断当前访问负载水平。",
            "P95 用于衡量大部分用户的真实响应体验。",
            "5xx 错误率用于判断服务正确性是否下降。",
            "服务可用性用于快速识别实例是否存活。",
            "JVM Heap 与 GC 指标用于发现 Java 服务内部压力。",
            "Redis 内存与 OPS 用于判断缓存层是否成为瓶颈。",
        ],
    )
    add_image(doc, "overview.png", "图 8-1 系统总览监控看板")

    doc.add_heading("9. 业务链路指标设计", level=1)
    doc.add_paragraph(
        "除了系统总览外，项目还设计了业务链路看板，用来回答“到底是哪条业务链路出了问题”。"
        "这一部分非常适合在毕业设计中突出，因为它体现了系统从技术监控走向业务监控的能力。"
    )
    add_bullets(
        doc,
        [
            "Service 调用速率：统计 admin 与 project 层各服务方法的调用频率。",
            "Service 方法 P95：衡量具体方法的长尾延迟。",
            "用户中心链路：关注注册、登录、登出等用户行为。",
            "短链接核心链路：关注创建、更新、配额消费等关键动作。",
            "回收站链路：关注移入、恢复、彻底删除等操作量。",
            "统计分析链路：关注单链接统计、分组统计、访问记录查询。",
            "HTTP 业务接口吞吐：按 URI 与状态码拆分业务接口表现。",
        ],
    )
    add_image(doc, "business-chain.png", "图 9-1 业务链路监控看板")

    doc.add_heading("10. 自动埋点与业务埋点设计", level=1)
    doc.add_paragraph(
        "为了让 Grafana 不只是展示通用 HTTP 指标，系统在代码层设计了两种埋点机制。第一种是自动埋点，第二种是业务埋点。"
    )
    add_bullets(
        doc,
        [
            "自动埋点由 ServiceMetricsAspect 完成，它统一拦截 admin.service.impl 与 project.service.impl 的方法调用，自动记录调用次数、成功率和延迟分布。",
            "业务埋点由 BizMetricsKit 完成，它让开发者可以在关键业务节点显式打点，例如短链接创建成功、冲突、配额耗尽、统计查询为空、回收站恢复成功等。",
            "自动埋点更适合看服务层性能，业务埋点更适合解释业务状态变化。",
        ],
    )
    doc.add_paragraph(
        "这两类埋点结合后，系统既能看到“哪个方法慢了”，也能看到“为什么慢、哪个业务状态变差了”，从而建立从系统症状到业务根因的诊断路径。"
    )

    doc.add_heading("11. 告警规则设计", level=1)
    doc.add_paragraph(
        "在告警设计上，系统并没有依赖 Grafana 面板人工观察，而是把告警规则写在 Prometheus Rule 中持续评估。"
        "当前规则主要覆盖以下几类："
    )
    add_bullets(
        doc,
        [
            "服务不可用：聚合服务 down 时触发 critical 告警。",
            "高错误率：5 分钟窗口 5xx 错误率超过 5% 时告警。",
            "高延迟：P95 连续超过 800ms 时告警。",
            "JVM 内存与 GC：堆使用率和平均 GC 停顿异常时告警。",
            "Redis 异常：Redis 不可用或内存占用过高时告警。",
            "SLO 错误预算：基于错误预算燃烧率区分 P1、P2、P3 等不同等级故障。",
        ],
    )
    doc.add_paragraph(
        "这种设计方式的优点是：告警逻辑是版本化和可审计的，不会因为面板临时修改而失去一致性。"
    )

    doc.add_heading("12. Grafana 设计的实际价值", level=1)
    doc.add_paragraph(
        "对短链接系统而言，Grafana 的价值不只是展示几个折线图，而是在出现问题时快速回答三个问题："
    )
    add_bullets(
        doc,
        [
            "系统现在是不是已经故障。",
            "故障主要影响的是哪条业务链路。",
            "问题更像是应用层、缓存层、数据库层还是基础资源层造成的。",
        ],
    )
    doc.add_paragraph(
        "例如，当跳转失败率上升时，可以先看 5xx 错误率与服务可用性，再看配额耗尽指标和业务调用错误率，最后结合日志与 Redis 指标定位到底是缓存问题、配额问题还是服务异常。"
    )

    doc.add_heading("13. 结论", level=1)
    doc.add_paragraph(
        "综上，短链接系统在高并发设计上并不是简单地“加缓存”，而是围绕创建与跳转两条核心链路，分别采用布隆过滤器、空值缓存、分布式锁、事务、原子 SQL 和异步统计等机制构建完整的并发治理方案。"
    )
    doc.add_paragraph(
        "在监控设计上，系统则通过 Prometheus 与 Grafana 构建了从基础资源指标到业务语义指标的分层可观测性体系，并通过自动埋点、业务埋点与告警规则，把系统运行状态转化为可量化、可解释、可处置的证据链。"
    )
    doc.add_paragraph(
        "如果用于毕业设计，这两部分内容非常适合分别写入“高并发场景下的系统设计”和“可观测性与监控指标设计”章节，用于突出项目的工程深度。"
    )

    if OUT_PATH.exists():
        OUT_PATH.unlink()
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
