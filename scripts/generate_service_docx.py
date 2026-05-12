from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


MD_PATH = Path("/Users/heybox/Desktop/毕业设计/毕业设计项目version1/shortlink/docs/shortlink-system-service-interface-overview.md")
OUT_PATH = Path("/Users/heybox/Desktop/毕业设计/毕业设计项目version1/shortlink/docs/shortlink-system-service-interface-overview.docx")


def clean_inline(text: str) -> str:
    return text.replace("`", "").strip()


def main() -> None:
    base_dir = MD_PATH.parent
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("#### "):
            doc.add_heading(clean_inline(stripped[5:]), level=4)
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=2)
            continue
        if stripped.startswith("# "):
            title = doc.add_heading(clean_inline(stripped[2:]), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            image_rel = stripped.split("](", 1)[1][:-1]
            image_path = (base_dir / image_rel).resolve()
            if image_path.exists():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(5.8))
            else:
                doc.add_paragraph(f"[缺失图片] {image_rel}")
            continue

        if raw.startswith("  - "):
            style_name = "List Bullet 2" if "List Bullet 2" in doc.styles else "List Bullet"
            paragraph = doc.add_paragraph(style=style_name)
            paragraph.add_run(clean_inline(raw.strip()[2:]))
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(clean_inline(stripped[2:]))
            continue

        paragraph = doc.add_paragraph(clean_inline(stripped))
        if stripped.startswith("图 "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if OUT_PATH.exists():
        OUT_PATH.unlink()
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
